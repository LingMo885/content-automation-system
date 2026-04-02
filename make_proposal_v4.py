#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
import datetime, os, base64

doc = Document()

section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_img_from_path(doc, img_path, width_inches):
    """Add image by path, handling encoding properly for PDF conversion"""
    if not os.path.exists(img_path):
        return
    # Use inline image for better PDF compatibility
    with open(img_path, 'rb') as f:
        img_data = f.read()
    # Create image part
    from docx.shared import Mm
    from docx.oxml import OxmlElement
    from docx.oxml.ns import nsmap
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    # Add image using add_picture (most reliable)
    run.add_picture(img_path, width=Inches(width_inches))

# ============ 封面 ============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('江苏武荣装备科技有限公司')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(31, 73, 125)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('JIANG SU WU RONG EQUIPMENT TECHNOLOGY CO., LTD.')
run2.font.size = Pt(10)
run2.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('电动大门采购方案')
run3.font.size = Pt(28)
run3.font.bold = True
run3.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph()

p_line = doc.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_line = p_line.add_run('━' * 40)
run_line.font.color.rgb = RGBColor(192, 0, 0)
run_line.font.size = Pt(10)

doc.add_paragraph()

p_price = doc.add_paragraph()
p_price.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_price = p_price.add_run('含税报价：¥ 45,000 元')
run_price.font.size = Pt(18)
run_price.font.bold = True
run_price.font.color.rgb = RGBColor(192, 0, 0)

p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_date = p_date.add_run(f'报价日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
run_date.font.size = Pt(11)
run_date.font.color.rgb = RGBColor(80, 80, 80)

p_client = doc.add_paragraph()
p_client.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_client = p_client.add_run('客户：中国人民武装警察部队扬州支队')
run_client.font.size = Pt(11)
run_client.font.color.rgb = RGBColor(80, 80, 80)

doc.add_page_break()

# ============ 第一部分：产品概述 ============
add_heading(doc, '一、产品概述', 1)

p = doc.add_paragraph()
p.add_run('我司提供整套电动大门系统解决方案，包含门扇、立柱、电机、悬浮配件及车辆识别系统，包安装调试。')

doc.add_paragraph()

# 产品规格表格
add_heading(doc, '产品规格明细', 2)
tbl = doc.add_table(rows=12, cols=2)
tbl.style = 'Table Grid'
specs = [
    ('门宽', '6950mm'),
    ('门高', '1300mm'),
    ('扇数', '6扇（3+3对开）'),
    ('款式', '竖条百叶（枫叶款）'),
    ('门扇颜色', '深灰色'),
    ('立柱', '150×150×3.0mm（2根）'),
    ('外框', '60×80×3.0mm（面80mm）'),
    ('竖条', '40×20×1.5mm'),
    ('锁板', '60×150×2.85mm'),
    ('电机', '鸿运品牌'),
    ('悬浮配件', '不锈钢材质'),
    ('控制系统', '车辆识别系统'),
]
for i, (k, v) in enumerate(specs):
    row = tbl.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    set_cell_bg(row.cells[0], 'DCE6F1')

doc.add_page_break()

# ============ 第二部分：产品图纸/效果图 ============
add_heading(doc, '二、产品效果图', 1)

p = doc.add_paragraph()
p.add_run('我司深灰色竖条百叶电动门效果图：')

doc.add_paragraph()

img1 = '/Users/yyf/.openclaw/media/inbound/ebe0bd48-110b-4a7b-98f0-f2c1bc662395.png'
add_img_from_path(doc, img1, 5.8)

doc.add_paragraph()

p2 = doc.add_paragraph()
p2.add_run('门扇细节图：')

img2 = '/Users/yyf/.openclaw/media/inbound/4af56a6f-4ea4-44b2-b693-f200fc438066.png'
add_img_from_path(doc, img2, 5.0)

doc.add_page_break()

# ============ 第三部分：价格清单 ============
add_heading(doc, '三、含税报价清单', 1)

tbl2 = doc.add_table(rows=8, cols=3)
tbl2.style = 'Table Grid'
headers = ['序号', '项目', '金额（元）']
for j, h in enumerate(headers):
    cell = tbl2.rows[0].cells[j]
    cell.text = h
    cell.paragraphs[0].runs[0].font.bold = True
    set_cell_bg(cell, '1F497D')
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

price_rows = [
    ('1', '门扇（6950×1300）', '12,468'),
    ('2', '立柱（150×150）× 2根', '900'),
    ('3', '鸿运电机', '4,000'),
    ('4', '不锈钢悬浮配件', '3,600'),
    ('5', '车辆识别系统', '4,500'),
    ('6', '安装调试费', '2,200'),
    ('', '含税总报价', '¥ 45,000'),
]
for i, (a, b, c) in enumerate(price_rows):
    row = tbl2.rows[i+1]
    row.cells[0].text = a
    row.cells[1].text = b
    row.cells[2].text = c
    if i == 5:  # 总报价行
        for cell in row.cells:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.bold = True
                set_cell_bg(cell, 'E2EFDA')
    if i == 6:
        for cell in row.cells:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(192, 0, 0)
                set_cell_bg(cell, 'FCE4D6')

doc.add_paragraph()

p_note = doc.add_paragraph()
run_note = p_note.add_run('注：以上报价为含税含票价，含运输及安装调试费用。')
run_note.font.size = Pt(10)
run_note.font.color.rgb = RGBColor(100, 100, 100)
run_note.font.italic = True

doc.add_paragraph()

# ============ 第四部分：商务条款 ============
add_heading(doc, '四、商务条款', 1)

tbl3 = doc.add_table(rows=3, cols=2)
tbl3.style = 'Table Grid'
terms = [
    ('供货周期', '合同签订后15个工作日完成制作安装'),
    ('付款方式', '3:3:4（定金30%/发货40%/验收30%）'),
    ('质保期', '门扇2年，电机1年，终身维护'),
]
for i, (k, v) in enumerate(terms):
    row = tbl3.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    set_cell_bg(row.cells[0], 'DCE6F1')

doc.add_paragraph()

# ============ 第五部分：联系信息 ============
add_heading(doc, '五、联系信息', 1)

tbl4 = doc.add_table(rows=4, cols=2)
tbl4.style = 'Table Grid'
contact_data = [
    ('公司名称', '江苏武荣装备科技有限公司'),
    ('联系人', '【请填写】'),
    ('联系电话', '【请填写】'),
    ('邮    箱', '【请填写】'),
]
for i, (k, v) in enumerate(contact_data):
    row = tbl4.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    set_cell_bg(row.cells[0], 'DCE6F1')

doc.add_paragraph()
doc.add_paragraph()

# 底部声明
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_footer = p_footer.add_run('本方案仅供贵方内部评审使用，未经授权不得对外传播。')
run_footer.font.size = Pt(9)
run_footer.font.color.rgb = RGBColor(128, 128, 128)
run_footer.font.italic = True

output_docx = '/Users/yyf/Desktop/【投标项目】/扬州支队项目/扬州支队大门方案_江苏武荣_v4.0.docx'
doc.save(output_docx)
print(f'DOCX已保存: {output_docx}')
