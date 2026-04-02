#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

doc = Document()

section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_img(doc, img_path, width=5.5):
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============ 封面 ============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('江苏武荣装备科技有限公司')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(31, 73, 125)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('JIANG SU WU RONG EQUIPMENT TECHNOLOGY CO., LTD.')
run2.font.size = Pt(10)
run2.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('办公家具采购方案')
run3.font.size = Pt(28)
run3.font.bold = True
run3.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph()

p_line = doc.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_line = p_line.add_run('━' * 40)
run_line.font.color.rgb = RGBColor(192, 0, 0)
run_line.font.size = Pt(10)

doc.add_paragraph()

p_price = doc.add_paragraph()
p_price.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_price = p_price.add_run('含税报价：¥ 11,300 元')
run_price.font.size = Pt(18)
run_price.font.bold = True
run_price.font.color.rgb = RGBColor(192, 0, 0)

p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_date = p_date.add_run(f'报价日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
run_date.font.size = Pt(11)
run_date.font.color.rgb = RGBColor(80, 80, 80)

p_client = doc.add_paragraph()
p_client.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_client = p_client.add_run('客户：中国人民武装警察部队江苏总队')
run_client.font.size = Pt(11)
run_client.font.color.rgb = RGBColor(80, 80, 80)

doc.add_page_break()

# ============ 产品概述 ============
add_heading(doc, '一、产品概述', 1)
p = doc.add_paragraph()
p.add_run('我司提供现代化办公家具解决方案，适用于领导办公室及办公区域配置。产品涵盖办公班台、办公椅及会客沙发区，简约大气，符合部队机关办公环境要求。')

doc.add_paragraph()

# 产品效果图
add_heading(doc, '二、产品效果图', 1)
p = doc.add_paragraph()
p.add_run('现代简约风领导办公室效果图：')

img1 = '/Users/yyf/.openclaw/media/inbound/48dd813f-4d42-4471-b254-2ecc6623fccc.png'
add_img(doc, img1, 5.5)

doc.add_paragraph()

p2 = doc.add_paragraph()
p2.add_run('办公室平面布局图（14.28㎡）：')

img2 = '/Users/yyf/.openclaw/media/inbound/8576a31f-a6bf-40e6-93e7-12c6f565bc18.png'
add_img(doc, img2, 5.0)

doc.add_page_break()

# ============ 产品清单 ============
add_heading(doc, '三、产品清单及报价', 1)

tbl = doc.add_table(rows=6, cols=5)
tbl.style = 'Table Grid'
headers = ['序号', '产品名称', '数量', '单价（元）', '金额（元）']
for j, h in enumerate(headers):
    cell = tbl.rows[0].cells[j]
    cell.text = h
    cell.paragraphs[0].runs[0].font.bold = True
    set_cell_bg(cell, '1F497D')
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

items = [
    ('1', '办公班台', '1', '2,200', '2,200'),
    ('2', '办公椅子', '6', '350', '2,100'),
    ('3', '6人位会客沙发区', '1', '3,600', '3,600'),
    ('4', '会客沙发配套椅', '6', '350', '2,100'),
    ('', '含税总计', '', '', '¥ 11,300'),
]
for i, row_data in enumerate(items):
    row = tbl.rows[i+1]
    for j, val in enumerate(row_data):
        row.cells[j].text = val
        if i == 4:  # 总计行
            row.cells[j].paragraphs[0].runs[0].font.bold = True
            if j == 4:
                row.cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(192, 0, 0)
                set_cell_bg(row.cells[j], 'FCE4D6')
            else:
                set_cell_bg(row.cells[j], 'E2EFDA')

doc.add_paragraph()

p_note = doc.add_paragraph()
run_note = p_note.add_run('注：以上报价为含税含票价，含运输及安装调试费用。')
run_note.font.size = Pt(10)
run_note.font.italic = True
run_note.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# ============ 产品规格 ============
add_heading(doc, '四、产品规格说明', 1)

tbl2 = doc.add_table(rows=4, cols=2)
tbl2.style = 'Table Grid'
specs = [
    ('办公班台', '标准尺寸，环保板材，深色木纹台面，金属脚架，简洁大气'),
    ('办公椅子', '人体工学设计，网布透气椅背，可调节扶手，尼龙五星脚轮'),
    ('6人位会客沙发区', '现代简约风格，白色/浅灰色布艺或皮质坐垫，配套茶几'),
    ('会客配套椅', '与沙发区风格统一，舒适耐用，适合办公接待'),
]
for i, (k, v) in enumerate(specs):
    row = tbl2.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    set_cell_bg(row.cells[0], 'DCE6F1')

doc.add_paragraph()

# ============ 商务条款 ============
add_heading(doc, '五、商务条款', 1)

tbl3 = doc.add_table(rows=3, cols=2)
tbl3.style = 'Table Grid'
terms = [
    ('供货周期', '合同签订后15个工作日'),
    ('付款方式', '3:3:4（定金30%/发货40%/验收30%）'),
    ('质保期', '1年（非人为损坏），终身维护'),
]
for i, (k, v) in enumerate(terms):
    row = tbl3.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    set_cell_bg(row.cells[0], 'DCE6F1')

doc.add_paragraph()

# ============ 联系信息 ============
add_heading(doc, '六、联系信息', 1)

tbl4 = doc.add_table(rows=4, cols=2)
tbl4.style = 'Table Grid'
contact_data = [
    ('公司名称', '江苏武荣装备科技有限公司'),
    ('联系人', '【请填写】'),
    ('联系电话', '【请填写】'),
    ('邮    箱', '【请填写】'),
]
for i, (k, v) in enumerate(contact_data):
    row = tbl4.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    set_cell_bg(row.cells[0], 'DCE6F1')

doc.add_paragraph()
doc.add_paragraph()

p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_footer = p_footer.add_run('本方案仅供贵方内部评审使用，未经授权不得对外传播。')
run_footer.font.size = Pt(9)
run_footer.font.italic = True
run_footer.font.color.rgb = RGBColor(128, 128, 128)

output = '/Users/yyf/Desktop/【投标项目】/江苏总队/江苏总队办公家具采购方案_江苏武荣_v1.0.docx'
os.makedirs(os.path.dirname(output), exist_ok=True)
doc.save(output)
print(f'DOCX已保存: {output}')
