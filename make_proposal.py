#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

# ============ 封面 ============
# 公司名
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('江苏武荣装备科技有限公司')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(31, 73, 125)

# 副标题
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('JIANG SU WU RONG EQUIPMENT TECHNOLOGY CO., LTD.')
run2.font.size = Pt(10)
run2.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# 方案标题
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('武警扬州支队大门更换方案')
run3.font.size = Pt(26)
run3.font.bold = True
run3.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph()

# 红色分隔线效果
p_line = doc.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_line = p_line.add_run('━' * 40)
run_line.font.color.rgb = RGBColor(192, 0, 0)
run_line.font.size = Pt(10)

doc.add_paragraph()

# 报价
p_price = doc.add_paragraph()
p_price.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_price = p_price.add_run('含税报价：¥ 45,000 元')
run_price.font.size = Pt(16)
run_price.font.bold = True
run_price.font.color.rgb = RGBColor(192, 0, 0)

# 日期
p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_date = p_date.add_run(f'报价日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
run_date.font.size = Pt(11)
run_date.font.color.rgb = RGBColor(80, 80, 80)

doc.add_page_break()

# ============ 第一部分：项目背景 ============
add_heading(doc, '一、项目背景', 1)

p = doc.add_paragraph()
p.add_run('扬州支队现用红门电动门（K300G+HP30）故障需更换。我司受邀提交替代方案，具体信息如下：')

# 表格：原门参数
tbl = doc.add_table(rows=6, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [
    ('型号', '红门 K300G + HP30'),
    ('门宽', '6950mm'),
    ('门高', '1300mm'),
    ('扇数', '6扇'),
    ('款式', '枫叶款竖条百叶'),
    ('原报价', '5万+'),
]
for i, (k, v) in enumerate(data):
    row = tbl.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    set_cell_bg(row.cells[0], 'DCE6F1')

doc.add_paragraph()

# ============ 第二部分：我司方案 ============
add_heading(doc, '二、我司方案', 1)

p = doc.add_paragraph()
p.add_run('全套含门扇、立柱、电机、悬浮配件、车辆识别系统，包安装。我司方案相比红门原方案具有以下优势：')

# 对比表格
tbl2 = doc.add_table(rows=5, cols=3)
tbl2.style = 'Table Grid'
headers = ['对比项', '红门原方案', '武荣替代方案']
for j, h in enumerate(headers):
    cell = tbl2.rows[0].cells[j]
    cell.text = h
    cell.paragraphs[0].runs[0].font.bold = True
    set_cell_bg(cell, '1F497D')
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

rows2 = [
    ('报价', '5万+', '¥ 4.5万'),
    ('门扇颜色', '红色（过于醒目）', '深灰色（军标色）'),
    ('主机', '需新配', '鸿运品牌（配套）'),
    ('工期', '定制周期长', '仅门扇定制'),
]
for i, (a, b, c) in enumerate(rows2):
    row = tbl2.rows[i+1]
    row.cells[0].text = a
    row.cells[1].text = b
    row.cells[2].text = c
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    set_cell_bg(row.cells[0], 'DCE6F1')

doc.add_paragraph()

# ============ 第三部分：产品规格 ============
add_heading(doc, '三、产品规格明细', 1)

tbl3 = doc.add_table(rows=12, cols=2)
tbl3.style = 'Table Grid'
specs = [
    ('门宽', '6950mm'),
    ('门高', '1300mm'),
    ('扇数', '6扇（3+3对开）'),
    ('款式', '竖条百叶（枫叶款）'),
    ('颜色', '深灰色（军标色）'),
    ('立柱', '150×150×3.0mm（2根）'),
    ('外框', '60×80×3.0mm（面80mm）'),
    ('竖条', '40×20×1.5mm'),
    ('锁板', '60×150×2.85mm'),
    ('电机', '鸿运品牌'),
    ('悬浮配件', '不锈钢材质'),
    ('控制系统', '车辆识别系统'),
]
for i, (k, v) in enumerate(specs):
    row = tbl3.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    set_cell_bg(row.cells[0], 'DCE6F1')

doc.add_page_break()

# ============ 第四部分：价格明细 ============
add_heading(doc, '四、价格明细（含税含票）', 1)

tbl4 = doc.add_table(rows=9, cols=4)
tbl4.style = 'Table Grid'
headers4 = ['序号', '项目', '数量', '金额（元）']
for j, h in enumerate(headers4):
    cell = tbl4.rows[0].cells[j]
    cell.text = h
    cell.paragraphs[0].runs[0].font.bold = True
    set_cell_bg(cell, '1F497D')
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

price_rows = [
    ('1', '门扇（6950×1300）', '1套', '12,468'),
    ('2', '立柱（150×150）', '2根', '900'),
    ('3', '鸿运电机', '1套', '4,000'),
    ('4', '不锈钢悬浮配件', '1套', '3,600'),
    ('5', '车辆识别系统', '1套', '4,500'),
    ('6', '安装费', '1项', '2,200'),
    ('', '合计', '', '27,668'),
    ('', '成交价', '', '¥ 27,200'),
]
for i, (a, b, c, d) in enumerate(price_rows):
    row = tbl4.rows[i+1]
    row.cells[0].text = a
    row.cells[1].text = b
    row.cells[2].text = c
    row.cells[3].text = d
    if i >= 6:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.bold = True
            set_cell_bg(cell, 'E2EFDA')

doc.add_paragraph()

# 毛利分析
tbl5 = doc.add_table(rows=4, cols=2)
tbl5.style = 'Table Grid'
margin_data = [
    ('含税报价', '¥ 45,000 元'),
    ('实际成本', '¥ 27,200 元'),
    ('毛利润', '¥ 17,800 元'),
    ('毛利率', '约 40%'),
]
for i, (k, v) in enumerate(margin_data):
    row = tbl5.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    set_cell_bg(row.cells[0], 'DCE6F1')
    if i == 0:
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(192, 0, 0)
        row.cells[1].paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()

# ============ 第五部分：供货与售后 ============
add_heading(doc, '五、供货周期', 1)
p = doc.add_paragraph()
p.add_run('合同签订后 ').font.bold = False
run = p.add_run('15个工作日')
run.bold = True
run.font.color.rgb = RGBColor(192, 0, 0)
p.add_run(' 内完成制作，现场安装 ')
run2 = p.add_run('1个工作日')
run2.bold = True
run2.font.color.rgb = RGBColor(192, 0, 0)
p.add_run('，合计 ')
run3 = p.add_run('16个工作日')
run3.bold = True
run3.font.color.rgb = RGBColor(192, 0, 0)
p.add_run('。')

add_heading(doc, '六、付款方式', 1)
pays = [
    '合同签订付 30% 定金（¥ 13,500）',
    '制作完成发货前付 40%（¥ 18,000）',
    '安装验收合格后付 30%（¥ 13,500）',
]
for pay in pays:
    doc.add_paragraph(pay, style='List Number')

add_heading(doc, '七、售后服务', 1)
services = [
    '门扇质保：2年（非人为损坏）',
    '电机质保：1年',
    '终身维护：提供配件供应和技术支持',
    '响应时间：24小时内响应，48小时内到场',
]
for svc in services:
    doc.add_paragraph(svc, style='List Bullet')

doc.add_paragraph()

# ============ 第八部分：联系信息 ============
add_heading(doc, '八、联系信息', 1)
tbl6 = doc.add_table(rows=4, cols=2)
tbl6.style = 'Table Grid'
contact_data = [('公司名称', '江苏武荣装备科技有限公司'), ('联系人', ''), ('联系电话', ''), ('邮    箱', '')]
for i, (k, v) in enumerate(contact_data):
    row = tbl6.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    set_cell_bg(row.cells[0], 'DCE6F1')

doc.add_paragraph()
doc.add_paragraph()

# 底部声明
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_footer = p_footer.add_run('本方案仅供贵方内部评审使用，未经授权不得对外传播。')
run_footer.font.size = Pt(9)
run_footer.font.color.rgb = RGBColor(128, 128, 128)
run_footer.font.italic = True

# 保存
output_path = '/Users/yyf/Desktop/【投标项目】/扬州支队项目/扬州支队大门更换方案_江苏武荣_v2.0.docx'
doc.save(output_path)
print(f'方案书已保存: {output_path}')
